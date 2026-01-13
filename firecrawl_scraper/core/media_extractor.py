#!/usr/bin/env python3
"""
Media Extractor - PDF/DOCX/Image Content Extraction

Extract content from various media formats:
- PDF documents with text and structure
- DOCX documents
- Images with optional OCR
- Media metadata extraction

Usage:
    from firecrawl_scraper.core.media_extractor import MediaExtractor

    extractor = MediaExtractor()
    content = await extractor.extract_pdf("https://example.com/doc.pdf")
"""

import asyncio
import logging
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import tempfile
import aiohttp

logger = logging.getLogger(__name__)

# Try to import PDF library
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    logger.warning("pypdf not installed. PDF extraction disabled. Install with: pip install pypdf")

# Try to import DOCX library
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. DOCX extraction disabled. Install with: pip install python-docx")

# Try to import PIL for images
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    logger.warning("Pillow not installed. Image processing disabled. Install with: pip install pillow")


@dataclass
class ExtractedMedia:
    """Extracted content from media file"""
    source_url: str
    media_type: str
    content: str
    metadata: Dict[str, Any]
    page_count: int = 0
    word_count: int = 0
    extraction_time: float = 0.0
    success: bool = True
    error: Optional[str] = None


class MediaExtractor:
    """
    Extract content from various media formats.

    Supports:
    - PDF documents
    - DOCX documents
    - Images (with optional OCR)
    """

    SUPPORTED_TYPES = {
        'pdf': ['.pdf', 'application/pdf'],
        'docx': ['.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'image': ['.png', '.jpg', '.jpeg', '.gif', '.webp', 'image/png', 'image/jpeg', 'image/gif', 'image/webp']
    }

    def __init__(
        self,
        max_size_mb: int = 50,
        enable_ocr: bool = False,
        timeout: int = 60
    ):
        """
        Initialize media extractor.

        Args:
            max_size_mb: Maximum file size to process
            enable_ocr: Enable OCR for images
            timeout: Download timeout in seconds
        """
        self.max_size_mb = max_size_mb
        self.enable_ocr = enable_ocr
        self.timeout = timeout

    def detect_media_type(self, url: str) -> Optional[str]:
        """
        Detect media type from URL.

        Args:
            url: URL to check

        Returns:
            Media type string ('pdf', 'docx', 'image') or None
        """
        url_lower = url.lower()

        # Check extension
        for media_type, extensions in self.SUPPORTED_TYPES.items():
            for ext in extensions:
                if ext.startswith('.') and url_lower.endswith(ext):
                    return media_type

        # Try to guess from URL
        mime_type, _ = mimetypes.guess_type(url)
        if mime_type:
            for media_type, types in self.SUPPORTED_TYPES.items():
                if mime_type in types:
                    return media_type

        return None

    async def download_file(self, url: str) -> Optional[bytes]:
        """
        Download file from URL.

        Args:
            url: URL to download

        Returns:
            File bytes or None if failed
        """
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(
                    url,
                    timeout=aiohttp.ClientTimeout(total=self.timeout)
                ) as response:
                    if response.status != 200:
                        logger.error(f"Failed to download {url}: HTTP {response.status}")
                        return None

                    # Check file size
                    content_length = response.headers.get('Content-Length')
                    if content_length:
                        size_mb = int(content_length) / (1024 * 1024)
                        if size_mb > self.max_size_mb:
                            logger.error(f"File too large: {size_mb:.1f}MB > {self.max_size_mb}MB")
                            return None

                    return await response.read()

        except Exception as e:
            logger.error(f"Download error for {url}: {e}")
            return None

    async def extract(self, url: str) -> ExtractedMedia:
        """
        Extract content from media URL.

        Auto-detects media type and uses appropriate extractor.

        Args:
            url: URL to media file

        Returns:
            ExtractedMedia object with content
        """
        start_time = datetime.now()

        # Detect media type
        media_type = self.detect_media_type(url)
        if not media_type:
            return ExtractedMedia(
                source_url=url,
                media_type='unknown',
                content='',
                metadata={},
                success=False,
                error='Unsupported or undetectable media type'
            )

        # Route to appropriate extractor
        if media_type == 'pdf':
            result = await self.extract_pdf(url)
        elif media_type == 'docx':
            result = await self.extract_docx(url)
        elif media_type == 'image':
            result = await self.extract_image(url)
        else:
            result = ExtractedMedia(
                source_url=url,
                media_type=media_type,
                content='',
                metadata={},
                success=False,
                error=f'No extractor for media type: {media_type}'
            )

        result.extraction_time = (datetime.now() - start_time).total_seconds()
        return result

    async def extract_pdf(self, url: str) -> ExtractedMedia:
        """
        Extract content from PDF document.

        Args:
            url: URL to PDF file

        Returns:
            ExtractedMedia with PDF content
        """
        if not HAS_PYPDF:
            return ExtractedMedia(
                source_url=url,
                media_type='pdf',
                content='',
                metadata={},
                success=False,
                error='pypdf not installed'
            )

        # Download file
        file_bytes = await self.download_file(url)
        if not file_bytes:
            return ExtractedMedia(
                source_url=url,
                media_type='pdf',
                content='',
                metadata={},
                success=False,
                error='Failed to download PDF'
            )

        try:
            # Save to temp file and extract
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
                f.write(file_bytes)
                temp_path = f.name

            reader = PdfReader(temp_path)

            # Extract text from all pages
            content_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content_parts.append(text)

            content = '\n\n'.join(content_parts)

            # Extract metadata
            metadata = {}
            if reader.metadata:
                metadata = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                    'producer': reader.metadata.get('/Producer', ''),
                    'creation_date': str(reader.metadata.get('/CreationDate', '')),
                }

            # Clean up temp file
            Path(temp_path).unlink(missing_ok=True)

            return ExtractedMedia(
                source_url=url,
                media_type='pdf',
                content=content,
                metadata=metadata,
                page_count=len(reader.pages),
                word_count=len(content.split()),
                success=True
            )

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ExtractedMedia(
                source_url=url,
                media_type='pdf',
                content='',
                metadata={},
                success=False,
                error=str(e)
            )

    async def extract_docx(self, url: str) -> ExtractedMedia:
        """
        Extract content from DOCX document.

        Args:
            url: URL to DOCX file

        Returns:
            ExtractedMedia with DOCX content
        """
        if not HAS_DOCX:
            return ExtractedMedia(
                source_url=url,
                media_type='docx',
                content='',
                metadata={},
                success=False,
                error='python-docx not installed'
            )

        # Download file
        file_bytes = await self.download_file(url)
        if not file_bytes:
            return ExtractedMedia(
                source_url=url,
                media_type='docx',
                content='',
                metadata={},
                success=False,
                error='Failed to download DOCX'
            )

        try:
            # Save to temp file and extract
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                f.write(file_bytes)
                temp_path = f.name

            doc = Document(temp_path)

            # Extract paragraphs
            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        content_parts.append(row_text)

            content = '\n\n'.join(content_parts)

            # Extract metadata from core properties
            metadata = {}
            if doc.core_properties:
                metadata = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'created': str(doc.core_properties.created or ''),
                    'modified': str(doc.core_properties.modified or ''),
                }

            # Clean up temp file
            Path(temp_path).unlink(missing_ok=True)

            return ExtractedMedia(
                source_url=url,
                media_type='docx',
                content=content,
                metadata=metadata,
                page_count=len(doc.paragraphs),
                word_count=len(content.split()),
                success=True
            )

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ExtractedMedia(
                source_url=url,
                media_type='docx',
                content='',
                metadata={},
                success=False,
                error=str(e)
            )

    async def extract_image(self, url: str) -> ExtractedMedia:
        """
        Extract metadata from image (and OCR if enabled).

        Args:
            url: URL to image file

        Returns:
            ExtractedMedia with image metadata
        """
        if not HAS_PIL:
            return ExtractedMedia(
                source_url=url,
                media_type='image',
                content='',
                metadata={},
                success=False,
                error='Pillow not installed'
            )

        # Download file
        file_bytes = await self.download_file(url)
        if not file_bytes:
            return ExtractedMedia(
                source_url=url,
                media_type='image',
                content='',
                metadata={},
                success=False,
                error='Failed to download image'
            )

        try:
            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
                f.write(file_bytes)
                temp_path = f.name

            # Open with PIL
            img = Image.open(temp_path)

            # Extract metadata
            metadata = {
                'format': img.format,
                'mode': img.mode,
                'size': f"{img.width}x{img.height}",
                'width': img.width,
                'height': img.height,
            }

            # Extract EXIF data if available
            if hasattr(img, '_getexif') and img._getexif():
                exif = img._getexif()
                metadata['exif'] = {str(k): str(v) for k, v in exif.items() if isinstance(v, (str, int, float))}

            content = f"Image: {img.format} {img.width}x{img.height} {img.mode}"

            # OCR would go here if enabled
            if self.enable_ocr:
                content += "\n[OCR not implemented - requires additional dependencies]"

            # Clean up
            img.close()
            Path(temp_path).unlink(missing_ok=True)

            return ExtractedMedia(
                source_url=url,
                media_type='image',
                content=content,
                metadata=metadata,
                success=True
            )

        except Exception as e:
            logger.error(f"Image extraction error: {e}")
            return ExtractedMedia(
                source_url=url,
                media_type='image',
                content='',
                metadata={},
                success=False,
                error=str(e)
            )

    async def extract_batch(
        self,
        urls: List[str],
        max_concurrent: int = 5
    ) -> List[ExtractedMedia]:
        """
        Extract content from multiple media URLs.

        Args:
            urls: List of media URLs
            max_concurrent: Maximum concurrent extractions

        Returns:
            List of ExtractedMedia objects
        """
        semaphore = asyncio.Semaphore(max_concurrent)

        async def extract_with_semaphore(url: str) -> ExtractedMedia:
            async with semaphore:
                return await self.extract(url)

        tasks = [extract_with_semaphore(url) for url in urls]
        return await asyncio.gather(*tasks)
